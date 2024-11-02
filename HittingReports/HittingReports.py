# Purpose here is to generate a .doc file for each player

import pandas as pd
import matplotlib.pyplot as plt
import os
from docx import Document
from docx.shared import Inches
from datetime import datetime, timedelta

# Load data
data = pd.read_csv(r"C:\Users\noahl\OneDrive\Documents\Git Repos\ReplayPlayerDev\MasterData\MasterBlastData.csv")

# Convert the Date column to datetime format
data['Date'] = pd.to_datetime(data['Date'])

# Calculate the date range for the previous month
today = datetime.today()
first_day_of_current_month = today.replace(day=1)
last_day_of_previous_month = first_day_of_current_month - timedelta(days=1)
first_day_of_previous_month = last_day_of_previous_month.replace(day=1)

# Filter data for the previous month
previous_month_data = data[(data['Date'] >= first_day_of_previous_month) & (data['Date'] <= last_day_of_previous_month)]

# Output folders for plots and reports
today_date = today.strftime('%Y-%m-%d')
plot_folder = os.path.join(r"C:\Users\noahl\OneDrive\Documents\Git Repos\ReplayPlayerDev\HittingReports\Plots", today_date)
report_folder = os.path.join(r"C:\Users\noahl\OneDrive\Documents\Git Repos\ReplayPlayerDev\HittingReports\GeneratedReports", today_date)
os.makedirs(plot_folder, exist_ok=True)
os.makedirs(report_folder, exist_ok=True)

# Loop through each player in the filtered dataset
for player in previous_month_data["Player_Name"].unique():
    # Filter data for the current player
    player_data = previous_month_data[previous_month_data["Player_Name"] == player]

    # Create a plot for the player (e.g., Bat Speed vs. Rotational Acceleration)
    plt.figure(figsize=(6, 4))
    plt.scatter(player_data["Bat_Speed"], player_data["Rotational_Acceleration"], color="blue")
    plt.title("Bat Speed vs Rotational Acceleration")
    plt.xlabel("Bat Speed")
    plt.ylabel("Rotational Acceleration")

    # Save the plot as an image file for the player
    plot_path = os.path.join(plot_folder, f"{player}_bat_speed_vs_rotation.png")
    plt.savefig(plot_path)
    plt.close()

    # Create a new Word document for the player
    doc = Document()

    # Add a title with player's name and the date range
    date_range_str = f"{first_day_of_previous_month.strftime('%Y-%m-%d')} to {last_day_of_previous_month.strftime('%Y-%m-%d')}"
    doc.add_heading(f'Report for {player} ({date_range_str})', level=1)

    # Add some stats or summary for the player
    avg_bat_speed = player_data["Bat_Speed"].mean()
    avg_rotation = player_data["Rotational_Acceleration"].mean()
    doc.add_paragraph(f"Average Bat Speed: {avg_bat_speed:.2f}")
    doc.add_paragraph(f"Average Rotational Acceleration: {avg_rotation:.2f}")

    # Insert the plot image
    doc.add_picture(plot_path, width=Inches(5.5))

    # Add a comments section
    doc.add_heading('Comments', level=2)
    doc.add_paragraph("Add comments here...")

    # Save the document with player's name and date in the file name
    report_path = os.path.join(report_folder, f"{player}_report_{today_date}.docx")
    doc.save(report_path)
    print(f"Report generated for {player} at {report_path}")
